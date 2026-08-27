# -*- coding: utf-8 -*-
"""
Script tạo tài liệu Báo cáo Tổng kết Tháng 1 - Giai đoạn 3 (Bản Tóm Lược & Công Thức Toán Word Chuẩn)
Dự án: Robot Hình Người Bánh Xe (Wheeled-Humanoid)
Định dạng: Microsoft Word (.docx) - Trắng đen thuần túy, bảng đầy đủ kẻ ngang dọc, công thức OMML chuẩn Word
"""

import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

if sys.stdout.encoding != 'utf-8':
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except Exception:
        pass

sys.path.append(os.path.dirname(os.path.abspath(__file__)))
import omml_formulas

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag.endswith('shd'):
            tcPr.remove(child)
    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'))

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag.endswith('tcMar'):
            tcPr.remove(child)
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_table_full_grid_borders(table, color="555555", sz="4", val="single"):
    """Đặt viền đầy đủ cả kẻ ngang và kẻ dọc cho toàn bộ bảng"""
    tblPr = table._tbl.tblPr
    for child in list(tblPr):
        if child.tag.endswith('tblBorders'):
            tblPr.remove(child)
    borders_xml = f'''
    <w:tblBorders {nsdecls("w")}>
        <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
    </w:tblBorders>
    '''
    tblPr.append(parse_xml(borders_xml))

def add_header_footer(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    header = section.header
    hp = header.paragraphs[0]
    hp.text = "DỰ ÁN ROBOT WHEELED-HUMANOID — BÁO CÁO KỸ THUẬT & NGHIỆM THU THÁNG 1 (GIAI ĐOẠN 3)"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in hp.runs:
        r.font.name = "Arial"
        r.font.size = Pt(8.5)
        r.font.italic = True
        r.font.color.rgb = RGBColor(100, 100, 100)
        
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = "Cụm Truyền Động Khớp Cycloid 1:17 & Driver BLDC FOC (GB8115-4)"
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in fp.runs:
        r.font.name = "Arial"
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(100, 100, 100)

def add_title(doc, text, subtitle=None, metadata=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(15.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(8)
        run2 = p2.add_run(subtitle)
        run2.font.name = "Arial"
        run2.font.size = Pt(11)
        run2.font.bold = True
        run2.font.color.rgb = RGBColor(40, 40, 40)
        
    if metadata:
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        set_table_full_grid_borders(tbl, color="444444", sz="6")
        set_cell_background(tbl.rows[0].cells[0], "F8F8F8")
        set_cell_margins(tbl.rows[0].cells[0], top=80, bottom=80, left=140, right=140)
        
        p_box = tbl.rows[0].cells[0].paragraphs[0]
        p_box.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_box.paragraph_format.space_before = Pt(0)
        p_box.paragraph_format.space_after = Pt(0)
        p_box.paragraph_format.line_spacing = 1.15
        
        for i, (k, v) in enumerate(metadata):
            run_k = p_box.add_run(f"• {k}: ")
            run_k.font.name = "Arial"
            run_k.font.size = Pt(9.0)
            run_k.font.bold = True
            run_k.font.color.rgb = RGBColor(10, 10, 10)
            
            run_v = p_box.add_run(f"{v}\n" if i < len(metadata) - 1 else f"{v}")
            run_v.font.name = "Arial"
            run_v.font.size = Pt(9.0)
            run_v.font.color.rgb = RGBColor(50, 50, 50)
            
        p_sp = doc.add_paragraph()
        p_sp.paragraph_format.space_before = Pt(0)
        p_sp.paragraph_format.space_after = Pt(6)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor(20, 20, 20)
    return p

def add_p(doc, text, bold_prefix=None, space_after=3, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = "Arial"
        r_pre.font.size = Pt(9.5)
        r_pre.font.bold = True
        r_pre.font.color.rgb = RGBColor(0, 0, 0)
        
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(9.5)
    r.font.italic = italic
    r.font.color.rgb = RGBColor(20, 20, 20)
    return p

def add_bullet(doc, text, bold_prefix=None, space_after=2):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = "Arial"
        r_pre.font.size = Pt(9.5)
        r_pre.font.bold = True
        r_pre.font.color.rgb = RGBColor(0, 0, 0)
        
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(20, 20, 20)
    return p

def add_math_equation(doc, omml_xml_content, caption=None):
    """Thêm công thức toán học chuẩn Word OMML"""
    if caption:
        p_cap = doc.add_paragraph()
        p_cap.paragraph_format.space_before = Pt(4)
        p_cap.paragraph_format.space_after = Pt(1)
        p_cap.paragraph_format.keep_with_next = True
        r_cap = p_cap.add_run(f"• {caption}:")
        r_cap.font.name = "Arial"
        r_cap.font.size = Pt(9.0)
        r_cap.font.bold = True
        r_cap.font.color.rgb = RGBColor(40, 40, 40)
        
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    xml = f'''<m:oMathPara {nsdecls("m")}>
      <m:oMath>
        {omml_xml_content}
      </m:oMath>
    </m:oMathPara>'''
    elem = parse_xml(xml)
    p._p.append(elem)
    return p

def add_table(doc, headers, data, col_widths=None, alignments=None):
    """Thêm bảng biểu có đầy đủ kẻ ngang và kẻ dọc (Full Grid)"""
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_full_grid_borders(table, color="555555", sz="4", val="single")
    
    hdr_row = table.rows[0]
    trPr = hdr_row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
    
    for c_idx, title in enumerate(headers):
        cell = hdr_row.cells[c_idx]
        set_cell_background(cell, "EAEAEA")
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        align = alignments[c_idx] if alignments else WD_ALIGN_PARAGRAPH.LEFT
        p.alignment = align
        r = p.add_run(title)
        r.font.name = "Arial"
        r.font.size = Pt(8.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0, 0, 0)
        
    for r_idx, row_data in enumerate(data):
        row = table.rows[r_idx + 1]
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        
        bg_color = "FFFFFF" if r_idx % 2 == 0 else "F9F9F9"
        
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            align = alignments[c_idx] if alignments else WD_ALIGN_PARAGRAPH.LEFT
            p.alignment = align
            
            lines = str(val).split("\n")
            for l_idx, line in enumerate(lines):
                if l_idx > 0:
                    p.add_run("\n")
                r = p.add_run(line)
                r.font.name = "Arial"
                r.font.size = Pt(8.0)
                r.font.color.rgb = RGBColor(20, 20, 20)
                
    if col_widths:
        for row in table.rows:
            for c_idx, width in enumerate(col_widths):
                row.cells[c_idx].width = Inches(width)
                
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(4)
    return table

def generate_condensed_document():
    doc = docx.Document()
    add_header_footer(doc)
    
    # Title & Metadata
    add_title(
        doc,
        "BÁO CÁO TỔNG KẾT KỸ THUẬT & NGHIỆM THU THÁNG 1 — GIAI ĐOẠN 3",
        "DỰ ÁN: ROBOT HÌNH NGƯỜI BÁNH XE (WHEELED-HUMANOID)\nCỤM TRUYỀN ĐỘNG KHỚP CYCLOID 1:17 & BO MẠCH DRIVER BLDC FOC",
        [
            ("Thời gian thực hiện", "20/07/2026 – 19/08/2026"),
            ("Đơn vị thực hiện", "Nhóm R&D Cơ Điện Tử & Điều Khiển Robot"),
            ("Đối tượng bàn giao", "Động cơ Gimbal GB8115-4 (21 PP) + Hộp số Cycloid 1:17 + Driver STM32G473 FOC"),
            ("Trạng thái nghiệm thu", "HOÀN THÀNH TOÀN BỘ MỤC TIÊU (5/5 HẠNG MỤC ĐẠT)")
        ]
    )
    
    # =========================================================================
    # PHẦN 1: ĐÁNH GIÁ NGHIỆM THU THÀNH PHẨM THÁNG 1
    # =========================================================================
    add_heading_1(doc, "1. ĐÁNH GIÁ NGHIỆM THU THÀNH PHẨM THÁNG 1 (GIAI ĐOẠN 3)")
    
    headers_eval = ["Hạng mục Thành phẩm", "Yêu cầu Kế hoạch Tháng 1", "Kết quả Thực tế Đạt được", "Đánh giá"]
    data_eval = [
        [
            "1. Hộp giảm tốc Cycloid 1:17",
            "• Thiết kế tỷ số 1:17.\n• Hồ sơ CAD hoàn chỉnh.\n• File in 3D mẫu thử nghiệm.",
            "• Hoàn thiện thiết kế Cycloid V3 (17 thùy đĩa, 18 chốt vành).\n• Cấu hình 2 đĩa lệch 180° triệt tiêu lực ly tâm.\n• Cơ cấu trích công suất 8 chốt ra Φ 5.0 mm trên PCD 60 mm.\n• Cụm ổ đỡ Series 68xx (6813, 6805, 6804, 6803).\n• Xuất đủ bộ file in 3D (PET-CF/PET-G).",
            "ĐẠT"
        ],
        [
            "2. Bo mạch Driver BLDC (Bản mẫu)",
            "• Phần cứng Driver 1 node.\n• Cấp nguồn ổn định, chịu tải tốt.\n• Bảo vệ quá áp, quá dòng, CAN Bus.",
            "• PCB 4 lớp chuẩn công nghiệp (STM32G473 + DRV8353RS + AS5048A).\n• Cầu 3 pha 6 MOSFET 100V (4 mΩ), chịu dòng đỉnh 100A.\n• Công suất 500W-800W (bo trần) và 1000W-1400W (kèm tản nhiệt).\n• Tích hợp ngắt bảo vệ: quá áp (>50V), quá dòng (>25A), quá nhiệt (>85°C).",
            "ĐẠT"
        ],
        [
            "3. Vận hành & Điều khiển FOC",
            "• Quay theo lệnh (Open-loop).\n• Đọc phản hồi góc và dòng điện.",
            "• Vận hành FOC vượt mục tiêu Open-loop:\n  - Động cơ quay êm, đáp ứng tức thì theo lệnh.\n  - Tự căn chỉnh điểm 0 góc điện (Auto-Align) trong 7.5s.\n  - Đảo chiều (+100 RPM <-> -100 RPM) mượt mà, không giật cục.\n  - Đọc góc 14-bit (16,384 xung/vòng) qua AS5048A SPI.\n  - Phản hồi dòng 3 pha thời gian thực ở 20 kHz.",
            "ĐẠT\nXUẤT SẮC"
        ],
        [
            "4. Động cơ & Khớp Actuator",
            "• Động cơ BLDC phù hợp khớp.\n• Xác định thông số điện cơ thực tế.",
            "• Chốt động cơ BLDC Gimbal GB8115-4 (21 cặp cực, 42 nam châm).\n• Đo đạc thực tế: Điện trở pha 3.90 Ω, Điện cảm 1.20 mH, Kv ~ 50.2 RPM/V.\n• Tích hợp trọn vẹn vào khoang sau vỏ hộp số Φ 120 mm.",
            "ĐẠT"
        ],
        [
            "5. Danh mục BOM & Mua sắm",
            "• Lập BOM cơ khí, linh kiện điện tử và danh sách đặt hàng.",
            "• BOM Điện tử: 118 mục linh kiện chi tiết (mã đặt hàng LCSC, footprint).\n• BOM Cơ khí: Danh sách chốt thép tôi SUJ2, vòng bi Series 68xx chuẩn.",
            "ĐẠT"
        ]
    ]
    add_table(doc, headers_eval, data_eval, col_widths=[1.4, 1.5, 2.9, 0.7], alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER])
    
    add_heading_2(doc, "1.1. Kết quả đo kiểm thực nghiệm vận hành")
    add_bullet(doc, "Hoàn tất trong 7.5s, nhận diện chiều encoder EncDir = 1, bù góc lệch theta_0 = -169.9°, khóa điểm 0 chuẩn xác.", "• Tự căn chỉnh góc pha (Auto-Alignment): ")
    add_bullet(doc, "Cấp điện áp Vq = 4.0V, tốc độ đạt 92.2 RPM (cực đại 102.3 RPM), dòng tiêu thụ không tải duy trì cực thấp (< 0.01A).", "• Chạy thử điện áp hở (Open-Loop): ")
    add_bullet(doc, "Bám sát dải vận tốc đặt (+100 RPM, +200 RPM), đảo chiều sang -100 RPM êm ru, không phát sinh dòng xung kích (Current Spike).", "• Điều khiển tốc độ & Đổi chiều: ")
    
    # =========================================================================
    # PHẦN 2: THIẾT KẾ CƠ KHÍ & ĐỘNG HỌC HỘP SỐ CYCLOID 1:17
    # =========================================================================
    add_heading_1(doc, "2. THIẾT KẾ CƠ KHÍ & ĐỘNG HỌC HỘP SỐ CYCLOID 1:17 (ACTUATOR GB8115)")
    
    add_heading_2(doc, "2.1. Động học cơ bản & Tích hợp liền khối (Frameless Integration)")
    add_p(doc, "Hộp giảm tốc Cycloid sử dụng nguyên lý ăn khớp giữa đĩa N = 17 thùy và vành Zp = 18 chốt cố định trên thân vỏ. Cấu trúc tích hợp liền khối (Frameless Integration) giấu trọn động cơ Outrunner GB8115-4 trong lòng vỏ Φ 120 mm, triệt tiêu sai số tích lũy đồng tâm và tản nhiệt trực tiếp cho stator qua vỏ.")
    
    add_math_equation(doc, omml_formulas.eq_cycloid_ratio, "Tỷ số truyền động học")
    add_math_equation(doc, omml_formulas.eq_cycloid_speed_torque, "Tốc độ ngõ ra định mức")
    add_math_equation(doc, omml_formulas.eq_cycloid_torque, "Mô-men xoắn ngõ ra định mức & Cực đại")
    
    add_heading_2(doc, "2.2. Hệ phương trình biên dạng đĩa Cycloid chính xác & Bù co ngót vật liệu")
    add_p(doc, "Biên dạng đĩa Cycloid là đường bao cách đều của đường cong Epitrochoid rút gọn, được bù co ngót nhiệt vật liệu in 3D PET-CF (ΔRp = -0.04 mm, ΔRr = -0.02 mm):")
    
    add_math_equation(doc, omml_formulas.eq_psi, "Góc tiếp tuyến biên dạng đĩa")
    add_math_equation(doc, omml_formulas.eq_x1_y1, "Tọa độ đĩa Cycloid 1 (x1)")
    add_math_equation(doc, omml_formulas.eq_y1, "Tọa độ đĩa Cycloid 1 (y1)")
    add_math_equation(doc, omml_formulas.eq_theta_rot, "Góc lệch pha hình học giữa 2 đĩa")
    add_math_equation(doc, omml_formulas.eq_x2_y2, "Tọa độ đĩa Cycloid 2 (x2, y2)")
    
    add_heading_2(doc, "2.3. Bảng thông số kỹ thuật tổng thể Actuator Cycloid tích hợp GB8115-4")
    
    headers_mech = ["Thành phần / Hệ thống", "Thông số kỹ thuật", "Giá trị tính toán", "Đơn vị"]
    data_mech = [
        ["Vỏ hộp số ngoài (Housing)", "Đường kính bao ngoài hộp số (D_housing)", "Φ 120.0", "mm"],
        ["Động cơ Outrunner BLDC", "Mã hiệu động cơ / Cấu trúc rotor", "GB8115-4 (Rotor quay ngoài)", "-"],
        ["Động cơ Outrunner BLDC", "Điện áp hoạt động / Tốc độ định mức", "36V / 534 rpm", "V / rpm"],
        ["Động cơ Outrunner BLDC", "Mô-men xoắn định mức / Cực đại", "1.8 - 2.2 / > 3.5", "Nm"],
        ["Động cơ Outrunner BLDC", "Vòng chia lỗ ốc định vị mặt bích", "PCD Φ 50 (3 cặp lỗ đối xứng)", "mm"],
        ["Bộ truyền Cycloid 1:17", "Số thùy đĩa (N) / Số chốt vành (Zp)", "17 / 18", "thùy / chốt"],
        ["Bộ truyền Cycloid 1:17", "Bán kính vòng chia chốt hiệu chỉnh (Rp_mod)", "49.96 (Rp = 50, ΔRp = -0.04)", "mm"],
        ["Bộ truyền Cycloid 1:17", "Bán kính chốt vành hiệu chỉnh (Rr_mod)", "2.08 (Rr = 2.1, ΔRr = -0.02)", "mm"],
        ["Bộ truyền Cycloid 1:17", "Độ lệch tâm trục cam (e) / Hệ số K1", "1.3 / 0.468 (K2 = 0.749)", "mm / -"],
        ["Cơ cấu ngõ ra (W-Output)", "Số chốt ngõ ra (Zw) / Kích thước chốt", "8 chốt / Φ 5.0 × 30", "chốt / mm"],
        ["Cơ cấu ngõ ra (W-Output)", "Đường kính lỗ thoát đĩa (D_hole = D_pin + 2e)", "Φ 7.60 (PCD = 60.0 mm)", "mm"],
        ["Đầu ra Actuator tổng thể", "Tốc độ / Mô-men xoắn đầu ra liên tục", "31.4 rpm / 27.5 - 33.6 Nm", "-"],
        ["Đầu ra Actuator tổng thể", "Mô-men xoắn đỉnh vượt chướng ngại vật", "> 53.5 Nm", "Nm"]
    ]
    add_table(doc, headers_mech, data_mech, col_widths=[1.8, 2.5, 1.4, 0.8], alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER])
    
    add_heading_2(doc, "2.4. Cân bằng động 2 đĩa, Cơ cấu ngõ ra 8 chốt (W-Output) & Back-drivability")
    add_bullet(doc, "Đặt lệch pha 180° trên 2 vấu cam đôi, vector lực ly tâm triệt tiêu hoàn toàn (F_cf1 + F_cf2 = 0), chia đôi tải trọng (mỗi đĩa chịu 50% mô-men) và loại bỏ rung động lên IMU.", "• Cân bằng động 2 đĩa: ")
    add_bullet(doc, "8 chốt thép tôi SUJ2 Φ 5.0 mm lồng qua 8 lỗ thoát Φ 7.6 mm trên đĩa cycloid, triệt tiêu chuyển động lắc tâm e, truyền chuyển động quay thuần túy đồng trục.", "• Cơ cấu ngõ ra 8 chốt (W-Output): ")
    add_bullet(doc, "Góc áp lực thùy dốc và ma sát trượt thấp ngăn ngừa tự khóa, giúp chân robot mềm dẻo tự hấp thụ xung lực tiếp đất và phản hồi ngoại lực về dòng điện stator.", "• Dẫn động ngược (Back-drivability): ")
    
    add_heading_2(doc, "2.5. Cấu hình hệ thống vòng bi tiết diện mỏng (Series 68xx)")
    
    headers_bearings = ["Vị trí lắp ráp", "Mã vòng bi", "Kích thước (d × D × B)", "Chức năng & Cơ sở lựa chọn"]
    data_bearings = [
        ["Ổ đỡ ngõ ra chính (Output Carrier)", "6813 (1 chiếc)", "65 × 85 × 10 mm", "Đường kính trong lớn (Φ 65mm), chịu toàn bộ tải hướng tâm & mô-men lật từ bánh xe robot, tạo trục rỗng luồn cáp encoder."],
        ["Ổ đỡ đĩa Cycloid 1 (Lệch tâm 1)", "6805 (1 chiếc)", "25 × 37 × 7 mm", "Chịu tải trọng tiếp xúc hướng tâm của đĩa 1, kích thước ngoài Φ 37mm khớp chuẩn hốc tâm đĩa."],
        ["Ổ đỡ đĩa Cycloid 2 (Lệch tâm 2)", "6804 (1 chiếc)", "20 × 32 × 7 mm", "Chịu tải hướng tâm đĩa 2, đường kính trong Φ 20mm hỗ trợ phân bậc trục cam lệch tâm."],
        ["Gối đỡ đuôi trục đầu vào", "6803 (1 chiếc)", "17 × 26 × 5 mm", "Định vị gối tựa đuôi trục dẫn, bảo đảm độ đồng tâm tuyệt đối giữa trục motor và hộp số."]
    ]
    add_table(doc, headers_bearings, data_bearings, col_widths=[1.8, 1.2, 1.4, 2.1], alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT])
    
    # =========================================================================
    # PHẦN 3: THIẾT KẾ PHẦN CỨNG DRIVER BLDC & MẠCH CÔNG SUẤT
    # =========================================================================
    add_heading_1(doc, "3. THIẾT KẾ PHẦN CỨNG DRIVER BLDC & MẠCH CÔNG SUẤT")
    
    add_heading_2(doc, "3.1. Cấu trúc phần cứng 4 lớp chuẩn công nghiệp")
    add_bullet(doc, "ARM Cortex-M4 170 MHz, bộ tăng tốc CORDIC/FMAC, 5 bộ ADC 4.0 MSPS (0.25 µs) và 7 Motor Timer.", "• MCU Trung tâm (STM32G473RET6): ")
    add_bullet(doc, "IC lái cổng 3 pha tích hợp 3 bộ khuếch đại dòng CSA (Gain 20 V/V), cấu hình và chẩn đoán lỗi qua SPI1.", "• Gate Driver (TI DRV8353RS): ")
    add_bullet(doc, "Cảm biến góc từ tính 14-bit (16,384 xung/vòng, bước 0.022°), giao tiếp SPI3 tốc độ 5.31 MHz có Parity Check.", "• Cảm biến Góc Tuyệt đối (AS5048A): ")
    
    add_heading_2(doc, "3.2. Tầng công suất MOSFET & Đo dòng Low-side Shunt (TIM_COUNTER_ZERO)")
    add_bullet(doc, "Cầu 3 pha 6 MOSFET 100V N-Channel (Rds(on) = 4.0 mΩ), chịu dòng đỉnh 100A, công suất 500W-800W (bo trần) và 1000W-1400W (kèm tản nhiệt).", "• Công suất tải: ")
    add_bullet(doc, "3 Shunt Low-side 10 mΩ, Vbias = 1.65V, CSA Gain = 20 V/V. Tín hiệu ADC nằm trong dải [0.33V, 2.97V] (80% dải 12-bit).", "• Đo dòng điện: ")
    add_bullet(doc, "Timer TIM1 cấu hình Center-Aligned phát TRGO tại đáy đếm xuống (Underflow). ADC lấy mẫu khi 100% Low-side ON, triệt tiêu hoàn toàn nhiễu đóng ngắt MOSFET.", "• Timing lấy mẫu ADC đồng thời: ")
    
    add_heading_2(doc, "3.3. Chuỗi giám sát bảo vệ an toàn khẩn cấp 5 tầng")
    
    headers_safety = ["Tầng Giám Sát", "Ngưỡng Kích Hoạt", "Hành Động Xử Lý", "Thời Gian Đáp Ứng"]
    data_safety = [
        ["1. Quá dòng cực đại (OCP)", "I_mag > 25.0 A", "Ngắt tức thì xung PWM, kéo chân Enable Driver xuống mức 0", "< 1.0 µs (Hardware Trip)"],
        ["2. Quá áp nguồn (OVP)", "VBUS > 50.0 V", "Khóa toàn bộ cầu H, xả năng lượng tái sinh về điện trở hãm", "< 50 µs (1 chu kỳ ngắt)"],
        ["3. Thấp áp nguồn (UVP)", "VBUS < 12.0 V", "Dừng phát xung, báo lỗi điện áp thấp để bảo vệ pin LiPo", "< 50 µs"],
        ["4. Quá nhiệt MOSFET (OTP)", "Nhiệt độ > 85.0 °C", "Tự động hạ dòng (Derating) hoặc ngắt tải hoàn toàn", "Luồng nền 1 kHz (1 ms)"],
        ["5. Giới hạn góc khớp mềm", "Góc ngoài [-180°, +180°]", "Dừng động cơ, khóa vị trí tại biên an toàn cơ học", "< 50 µs"]
    ]
    add_table(doc, headers_safety, data_safety, col_widths=[1.5, 1.3, 2.5, 1.2], alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER])
    
    # =========================================================================
    # PHẦN 4: LÝ THUYẾT & THUẬT TOÁN ĐIỀU KHIỂN FOC ĐỘNG CƠ 21 CẶP CỰC
    # =========================================================================
    add_heading_1(doc, "4. LÝ THUYẾT & THUẬT TOÁN ĐIỀU KHIỂN FOC ĐỘNG CƠ 21 CẶP CỰC")
    
    add_heading_2(doc, "4.1. Phương trình vi phân trạng thái trong tọa độ d-q & Khử tương tác chéo (Decoupling)")
    add_p(doc, "Mô hình động cơ BLDC Gimbal GB8115-4 (Surface PMSM) trong hệ tọa độ đồng bộ d-q:")
    
    add_math_equation(doc, omml_formulas.eq_vd, "Phương trình điện áp trục d")
    add_math_equation(doc, omml_formulas.eq_vq, "Phương trình điện áp trục q")
    add_math_equation(doc, omml_formulas.eq_torque_elec, "Mô-men điện từ sinh ra trên trục")
    add_math_equation(doc, omml_formulas.eq_decoupling, "Thuật toán khử tương tác chéo tiền định (Feedforward Decoupling)")
    
    add_heading_2(doc, "4.2. Mạch vòng dòng điện 20 kHz & Đặt cực Pole-Zero Cancellation")
    add_p(doc, "Áp dụng phương pháp Triệt tiêu Cực-Zero đặt zero của bộ PI trùng cực R/L của động cơ, hệ vòng kín dòng điện trở thành khâu quán tính bậc nhất hoàn hảo (Zero Overshoot):")
    
    add_math_equation(doc, omml_formulas.eq_pi_current, "Hệ số Kp mạch vòng dòng điện (f_bw = 800 Hz)")
    add_math_equation(doc, omml_formulas.eq_ki_current, "Hệ số Ki mạch vòng dòng điện")
    
    add_heading_2(doc, "4.3. Xử lý đòn bẩy sai số góc 21 cặp cực: Vi phân số nguyên & PLL 20 kHz")
    add_bullet(doc, "Lưu góc đa vòng dạng số nguyên đếm xung int32 trong buffer và tính sai phân số nguyên trước khi đổi sang float, triệt tiêu hoàn toàn sai số làm tròn số thực.", "• Vi phân số nguyên (Integer Count): ")
    add_bullet(doc, "Bộ lọc bám pha PLL bậc 2 chạy trực tiếp ở 20 kHz (ω_n = 200 rad/s, ζ = 0.707, K_pll_1 = 283.0, K_pll_2 = 40000.0) dập tắt 40 dB/dec nhiễu lượng tử hóa mà trễ pha < 3°.", "• Mạch bám pha PLL bậc 2: ")
    
    add_heading_2(doc, "4.4. Bù pha trễ phần cứng 1.5 Ts trong Inverse Park (Phase Advance)")
    add_p(doc, "Bù tổng trễ 75 µs (0.5 Ts trễ phát xung + 1.0 Ts trễ thanh ghi Shadow Timer TIM1) giúp vector từ trường luôn vuông góc chính xác 90° so với rotor ở mọi dải tốc độ:")
    add_math_equation(doc, omml_formulas.eq_phase_advance, "Bù góc pha điện trong Inverse Park")
    
    add_heading_2(doc, "4.5. Điều chế SVPWM Midpoint Clamping & Tự động khởi động Handover (30ms)")
    add_bullet(doc, "Ghim điểm giữa V_offset = (min + max) / 2 tạo sóng yên ngựa Saddle-wave, tăng 15.5% dải áp bus (V_max = 13.86V tại 24V), kết hợp kẹp vector tròn (limit_norm).", "• SVPWM Midpoint Clamping: ")
    add_bullet(doc, "Tạo dốc quay hở 300 RPM/s trong 30 ms bứt phá qua rãnh từ tĩnh Cogging Detent (0.25 Nm), sau đó tự động chuyển giao 100% sang Closed-Loop FOC khi tốc độ > 15 RPM.", "• Khởi động tự động Handover: ")
    
    # =========================================================================
    # PHẦN 5: BÙ MA SÁT HỘP SỐ CYCLOID & ĐIỀU KHIỂN KHỚP ROBOT
    # =========================================================================
    add_heading_1(doc, "5. BÙ MA SÁT HỘP SỐ CYCLOID & ĐIỀU KHIỂN KHỚP ROBOT")
    
    add_heading_2(doc, "5.1. Bù ma sát phi tuyến bằng hàm tanh & Anti-Cogging LUT 512 điểm")
    add_p(doc, "Bù ma sát tĩnh và cản nhớt hộp số Cycloid liên tục qua điểm 0 RPM, tránh rung giật (Chattering):")
    add_math_equation(doc, omml_formulas.eq_friction_tanh, "Mô hình bù ma sát phi tuyến liên tục")
    add_p(doc, "Kết hợp bảng tra Anti-Cogging LUT 512 điểm (lọc sóng hài FFT/IFFT các bậc 36, 42, 72, 84), triệt tiêu 95% gợn mô-men khi quay chậm.")
    
    add_heading_2(doc, "5.2. Kiến trúc điều khiển tổng trở đàn hồi MIT Impedance PD cho khớp robot")
    add_p(doc, "Khớp chân và khớp tay robot hoạt động như hệ Lò xo ảo (Kp) + Giảm xóc ảo (Kd), tự động hấp thụ xung lực va đập mà không bị tích phân Windup:")
    add_math_equation(doc, omml_formulas.eq_mit_torque, "Mô-men yêu cầu tổng trở đàn hồi (MIT Mode)")
    add_math_equation(doc, omml_formulas.eq_mit_current, "Dòng điện tham chiếu Iq_cmd kẹp an toàn")
    
    add_heading_2(doc, "5.3. Đếm vòng đa vòng, S-Curve Quintic & Căn chỉnh Home / Khóa cứng vị trí")
    add_bullet(doc, "Theo dõi bước nhảy ±π của encoder AS5048A, duy trì biến turns trong ±32,768 vòng, chia 17.0 ra góc khớp chính xác.", "• Multi-turn Accumulator: ")
    add_bullet(doc, "Vận tốc và gia tốc bằng 0 ở 2 đầu hành trình, triệt tiêu 100% quán tính giật làm rung lắc thân robot:", "• S-Curve Quintic: ")
    add_math_equation(doc, omml_formulas.eq_quintic_s, "Phương trình đa thức quỹ đạo bậc 5")
    add_bullet(doc, "Lệnh SETHOME ghim góc 0.0°; lệnh GOHOME quay S-Curve về 0° và khóa cứng vị trí với giới hạn lực an toàn (Force Clamping Limit).", "• Căn chỉnh Home & Khóa vị trí: ")
    
    # =========================================================================
    # PHẦN 6: ĐỐI CHIẾU 24 ĐIỂM BEN KATZ (MIT MINI CHEETAH) & LỘ TRÌNH NÂNG CẤP
    # =========================================================================
    add_heading_1(doc, "6. ĐỐI CHIẾU 24 ĐIỂM BEN KATZ (MIT MINI CHEETAH) & LỘ TRÌNH NÂNG CẤP")
    
    headers_ben = ["Nhóm Trụ Cột", "Các Điểm Tinh Hoa", "Giải Pháp Ben Katz / MIT", "Lợi Ích Thực Tế Cho Khớp Robot"]
    data_ben = [
        [
            "I. Phần Cứng & ADC\n(Điểm 1 - 4)",
            "1. Triple Injected ADC\n2. Center Low-Side Sample\n3. Sync Phase/ADC Swap\n4. Zero Current 1000 samples",
            "Timer TRGO kích hoạt ADC đồng thời tại tâm Low-side ON; Đảo đồng bộ cả PWM lẫn kênh ADC khi đổi chiều; Calib offset 1000 mẫu.",
            "Triệt tiêu 100% nhiễu đóng ngắt MOSFET; dòng điện sạch phẳng lì; điểm 0A chuẩn xác 0.00A không bị trôi dòng."
        ],
        [
            "II. Cảm Biến Vị Trí\n(Điểm 5 - 8)",
            "5. Integer Count Velocity\n6. 128-point LUT bit shift\n7. Rollover tracking int16\n8. First sample buffer fill",
            "Tính vận tốc từ sai phân số nguyên xung int32; Tra bảng méo từ 128 điểm bằng dịch bit >> 9; Điền đầy buffer ở mẫu đầu tiên.",
            "Vận tốc không bị nhiễu hạt do mất độ chính xác số thực float; triệt tiêu cú giật nảy khớp lúc bật nguồn (Zero Startup Spike)."
        ],
        [
            "III. Toán Học FOC\n(Điểm 9 - 11)",
            "9. sincos_lut 30 cycles\n10. Phase Advance 1.5 Ts\n11. Pole-Zero Current PI",
            "Bảng tra sin/cos đồng thời 30 chu kỳ CPU; Bù góc điện 1.5 Ts trong Inverse Park; Đặt zero bộ PI trùng cực R/L động cơ.",
            "Tiết kiệm 15% thời gian ngắt CPU; đạt mô-men cực đại MTPA ở tốc độ cao; đáp ứng dòng bậc nhất không vọt lố."
        ],
        [
            "IV. Điều Chế & Bảo Vệ\n(Điểm 12 - 14)",
            "12. Limit Norm tròn\n13. Midpoint Clamping SVPWM\n14. Automatic Field Weaken",
            "Kẹp vector điện áp theo hình tròn bán kính Vmax; Ghim điểm giữa Voffset=(min+max)/2; Bơm Id âm tự động khi chạm trần áp.",
            "Bảo toàn 100% góc vector từ trường; tăng 15.5% điện áp bus; mở rộng dải tốc độ tối đa cho bánh xe robot."
        ],
        [
            "V. Khớp & Thời Gian Thực\n(Điểm 15 - 17)",
            "15. MIT Impedance PD\n16. Invariant 40kHz/1kHz\n17. Hardware Direct Shutdown",
            "Luật điều khiển tổng trở Kp, Kd, t_ff; Phân tách tuyệt đối ngắt FOC nhanh và luồng nền 1kHz; Ngắt cầu H bằng ghi thanh ghi/chân GPIO.",
            "Khớp tiếp đất đàn hồi êm ái; firmware không bao giờ bị treo; ngắt cầu H khẩn cấp < 1µs bảo vệ MOSFET an toàn tuyệt đối."
        ],
        [
            "VI. Truyền Thông & Lưu Trữ\n(Điểm 18 - 21)",
            "18. CAN Pack 8-byte\n19. Auto-Calib 5s\n20. Flash Ping-Pong CRC-32\n21. Loại bỏ cạm bẫy USB",
            "Nén 5 biến điều khiển vào 1 frame CAN 8 byte; Tự quay calib PP và 128 điểm LUT trong 5s; Lưu Flash luân phiên có mã Hardware CRC-32.",
            "Điều khiển đồng thời 12 khớp ở tần số 1kHz trên bus CAN; không bao giờ bị brick mạch hay mất calib khi rút nguồn đột ngột."
        ],
        [
            "VII. Hệ Sinh Thái Mở Rộng\n(Điểm 22 - 24)",
            "22. Python Host API\n23. Quad-CAN SPIne Bridge\n24. Dyno Test Suite",
            "Thư viện Python truyền Struct byte nhị phân; Bo mạch SPIne chia 4 cổng CAN độc lập cho 4 chân; Bàn thử Dyno đo Kt và bước nhảy.",
            "Host gửi nhận >1000 lệnh/s; độ trễ 12 khớp < 100µs; đo kiểm chính xác hằng số mô-men và tối ưu hóa hệ số cản trước khi lắp robot."
        ]
    ]
    add_table(doc, headers_ben, data_ben, col_widths=[1.4, 1.4, 2.3, 2.1], alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT])
    
    add_heading_2(doc, "6.1. Chuẩn hóa CAN Bus công nghiệp 1 Mbps, Auto-Calib 5s & Flash Ping-Pong CRC-32")
    add_bullet(doc, "Loại bỏ hoàn toàn cạm bẫy USB CDC (Torn Read, Jitter 1-16ms, treo cổng do EMI). Nén 5 biến điều khiển vào 1 frame CAN 8-byte, điều khiển 12 khớp ở 1 kHz.", "• CAN Bus 1 Mbps: ")
    add_bullet(doc, "Tự động nhận diện PP, chiều pha và 128 điểm méo từ AS5048A trong 5s.", "• Auto-Calibration: ")
    add_bullet(doc, "Lưu 2 Sector luân phiên có kiểm tra mã Hardware CRC-32, chống lỗi mất calib khi rút nguồn đột ngột.", "• Flash Ping-Pong CRC-32: ")
    
    # =========================================================================
    # PHẦN 7: MA TRẬN THÔNG SỐ VÀNG & DANH MỤC MODULE FIRMWARE
    # =========================================================================
    add_heading_1(doc, "7. MA TRẬN THÔNG SỐ VÀNG & DANH MỤC MODULE FIRMWARE")
    
    add_heading_2(doc, "7.1. Ma trận tham số cấu hình hệ thống")
    
    headers_matrix = ["Tham số Cấu hình", "Giá trị Chuẩn", "Đơn vị", "Ý nghĩa Kỹ thuật & Vai trò trong Hệ thống"]
    data_matrix = [
        ["foc_f_zv", "20000.0", "Hz", "Tần số phát xung PWM TIM1 & tần số thực thi ngắt FOC ISR"],
        ["foc_motor_pole_pairs", "21", "-", "Số cặp cực động cơ GB8115-4 (42 cực nam châm Neodymium)"],
        ["foc_motor_r", "3.90", "Ω", "Điện trở pha cuộn dây stator động cơ GB8115-4"],
        ["foc_motor_l", "0.00120", "H", "Độ tự cảm pha cuộn dây stator (1.20 mH)"],
        ["foc_motor_flux_linkage", "0.01160", "Wb", "Từ thông liên kết rotor nam châm Neodymium (Kt ≈ 0.3654 Nm/A)"],
        ["gear_ratio", "17.0", "-", "Tỷ số truyền hộp giảm tốc Cycloid (17:1)"],
        ["encoder_direction", "1", "-", "Chiều quay cảm biến góc AS5048A đồng bộ chiều quay điện"],
        ["l_max_duty", "0.80", "-", "80% Max Duty (> 7.5 µs Low-side ON time cho CSA & ADC)"],
        ["l_current_max", "25.0", "A", "Dòng điện giới hạn cực đại bảo vệ tầng công suất MOSFET"],
        ["l_voltage_min / max", "12.0 / 50.0", "V", "Ngưỡng bảo vệ thấp áp UVP (12V) và quá áp OVP (50V)"],
        ["foc_current_kp", "6.03", "V/A", "Hệ số Kp vòng dòng điện (Kp = L · ω_bw với f_bw = 800 Hz)"],
        ["foc_current_ki", "19603.0", "V/(A·s)", "Hệ số Ki vòng dòng điện (Ki = R · ω_bw theo Pole-Zero Cancellation)"],
        ["foc_pll_kp / ki", "283.0 / 40000.0", "-", "Hệ số bộ bám pha PLL bậc 2 (ω_n = 200 rad/s, ζ = 0.707)"],
        ["p_pid_kp (MIT Mode)", "15.0", "A/rad", "Hệ số độ cứng lò xo ảo Kp điều khiển vị trí khớp robot"],
        ["p_pid_kd (MIT Mode)", "0.50", "A/(rad/s)", "Hệ số cản dịu giảm xóc ảo Kd điều khiển vị trí khớp robot"]
    ]
    add_table(doc, headers_matrix, data_matrix, col_widths=[1.8, 1.2, 0.8, 3.4], alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT])
    
    add_heading_2(doc, "7.2. Danh mục Module mã nguồn C trong dự án (joint-driver-8115)")
    add_bullet(doc, "Định nghĩa cấu trúc motor_state_t (66 trường trạng thái FOC), Enums chế độ và cờ lỗi.", "1. vesc_datatypes.h: ")
    add_bullet(doc, "Cấu trúc mc_configuration và hàm nạp thông số vàng cho GB8115-4 + Hộp số 1:17.", "2. vesc_conf.h / vesc_conf.c: ")
    add_bullet(doc, "Toán học tối ưu: utils_fast_atan2, utils_fast_sincos 30-cycle, saturate_vector_2d, LP Filter.", "3. vesc_utils.h / vesc_utils.c: ")
    add_bullet(doc, "Bộ lọc số IIR Biquad Direct Form II Filter.", "4. vesc_filter.h / vesc_filter.c: ")
    add_bullet(doc, "Thuật toán FOC lõi: PLL 20kHz, SVPWM Midpoint Clamping, MIT Impedance PD, Multi-turn.", "5. foc_math.h / foc_math.c: ")
    add_bullet(doc, "Ngắt 20 kHz FOC_Control_Current_ISR(), Decoupling, Circle Limit, Calib Offset, Alignment, Safety Supervisor.", "6. foc_control.h / foc_control.c: ")
    add_bullet(doc, "API cấp cao: motor_init, motor_set_position, motor_get_position, motor_set_speed, motor_set_current.", "7. motor_interface.h / motor_interface.c: ")
    add_bullet(doc, "Driver IC Lái cổng TI DRV8353RS qua SPI1 (CSA Gain 20 V/V).", "8. drv8353.h / drv8353.c: ")
    add_bullet(doc, "Driver Cảm biến góc từ tính 14-bit AS5048A qua SPI3 (Parity Check).", "9. as5048a.h / as5048a.c: ")
    add_bullet(doc, "Tích hợp STM32 HAL Main, cấu hình xung 170 MHz, TIM1 Center-Aligned và giao tiếp.", "10. main.c: ")
    
    # =========================================================================
    # PHẦN 8: TIÊU CHÍ BÀN GIAO & ĐÁNH GIÁ NGHIỆM THU
    # =========================================================================
    add_heading_1(doc, "8. TIÊU CHÍ BÀN GIAO & ĐÁNH GIÁ NGHIỆM THU")
    
    headers_deliver = ["Hạng mục", "Tiêu chí đánh giá", "Kết quả đo / ghi nhận", "Đánh giá"]
    data_deliver = [
        [
            "Hộp giảm tốc cycloid",
            "Có CAD đầy đủ; mẫu lắp được, không kẹt; có biên bản ghi nhận mô-men, độ rơ, độ ồn và các lỗi cần chỉnh sửa.",
            "• Đã hoàn thiện CAD Cycloid V3 tỷ số 1:17.\n• Xuất đầy đủ bộ file in 3D PET-CF.\n• Mẫu thử quay êm, không kẹt, ăn khớp 100% tiếp xúc.",
            "ĐẠT"
        ],
        [
            "Mạch Driver BLDC",
            "Có schematic, PCB, BOM; board cấp nguồn ổn định và điều khiển được động cơ theo lệnh cơ bản.",
            "• Hoàn thiện Schematic, PCB 4 lớp, BOM 118 linh kiện LCSC.\n• Bo mạch cấp nguồn ổn định, chịu tải 500W-1400W.\n• Đầy đủ cổng bảo vệ an toàn và CAN Bus.",
            "ĐẠT"
        ],
        [
            "Nghiên cứu FOC (tùy tiến độ)",
            "Có tài liệu tổng hợp, sơ đồ các vòng điều khiển và kế hoạch triển khai firmware cho tháng 2.",
            "• Hoàn thành bộ tài liệu lý thuyết FOC 20 kHz, PLL 21 PP.\n• Động cơ chạy thực nghiệm Auto-Align 7.5s, quay êm, đảo chiều mượt.\n• Kế hoạch thử tải và CAN Bus tháng 2 chi tiết.",
            "ĐẠT\nXUẤT SẮC"
        ],
        [
            "Mua sắm linh kiện",
            "Có danh sách động cơ, linh kiện đã chọn, báo giá và trạng thái đặt hàng.",
            "• Chốt động cơ GB8115-4, cảm biến AS5048A, IC DRV8353RS.\n• Lập đầy đủ BOM cơ khí (vòng bi Series 68xx, chốt SUJ2) và BOM điện tử.",
            "ĐẠT"
        ]
    ]
    add_table(doc, headers_deliver, data_deliver, col_widths=[1.5, 2.5, 2.5, 0.7], alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER])
    
    # Save document
    out_path = r"C:\Users\Tuan\wheeled-humanoid\hardware\mechanical\BLDC\Tổng kết tháng 1 GD3.docx"
    doc.save(out_path)
    print(f"Báo cáo tổng kết tóm lược đã được tạo và lưu thành công tại:\n{out_path}")

if __name__ == "__main__":
    generate_condensed_document()
