import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_omml_para(doc, omml_xml_content):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    xml = f'''<m:oMathPara {nsdecls("m")}>
      <m:oMath>
        {omml_xml_content}
      </m:oMath>
    </m:oMathPara>'''
    elem = parse_xml(xml)
    p._p.append(elem)
    return p

def set_table_full_grid_borders(table, color="555555", sz="4"):
    """Đặt viền đầy đủ ngang và dọc cho toàn bộ bảng"""
    tblPr = table._tbl.tblPr
    for child in list(tblPr):
        if child.tag.endswith('tblBorders'):
            tblPr.remove(child)
    borders_xml = f'''
    <w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:insideV w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>
    </w:tblBorders>
    '''
    tblPr.append(parse_xml(borders_xml))

doc = docx.Document()
# Test OMML fraction and sub
math_content = '''
<m:r><m:t>i = </m:t></m:r>
<m:f>
  <m:num><m:sSub><m:e><m:r><m:t>n</m:t></m:r></m:e><m:sub><m:r><m:t>in</m:t></m:r></m:sub></m:sSub></m:num>
  <m:den><m:sSub><m:e><m:r><m:t>n</m:t></m:r></m:e><m:sub><m:r><m:t>out</m:t></m:r></m:sub></m:sSub></m:den>
</m:f>
<m:r><m:t> = </m:t></m:r>
<m:f>
  <m:num><m:r><m:t>N</m:t></m:r></m:num>
  <m:den><m:r><m:t>Z</m:t></m:r><m:sSub><m:e><m:r><m:t></m:t></m:r></m:e><m:sub><m:r><m:t>p</m:t></m:r></m:sub></m:sSub><m:r><m:t> - N</m:t></m:r></m:den>
</m:f>
<m:r><m:t> = </m:t></m:r>
<m:f>
  <m:num><m:r><m:t>17</m:t></m:r></m:num>
  <m:den><m:r><m:t>18 - 17</m:t></m:r></m:den>
</m:f>
<m:r><m:t> = 17:1</m:t></m:r>
'''
create_omml_para(doc, math_content)

# Test Table with full grid
tbl = doc.add_table(rows=2, cols=2)
set_table_full_grid_borders(tbl)
tbl.rows[0].cells[0].text = "H1"
tbl.rows[0].cells[1].text = "H2"
tbl.rows[1].cells[0].text = "D1"
tbl.rows[1].cells[1].text = "D2"

doc.save(r"C:\Users\Tuan\wheeled-humanoid\hardware\mechanical\BLDC\test_math.docx")
print("Saved test_math.docx successfully!")
